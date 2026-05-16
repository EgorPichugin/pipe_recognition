from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image
from pydantic import ValidationError

from src.core.config import _PROJECT_ROOT
from src.models.report import (
    FilteredReportChange,
    InspectionLocationPatch,
    ReportChangeRequest,
    ReportEditPlan,
    ReportIssue,
    ReportIssuePatch,
    ReportMetadataPatch,
    ReportNextStepPatch,
    ReportPartyPatch,
    ReportPatch,
    ReportSection,
    ReportSectionPatch,
    ReportSchema,
    build_dummy_report,
)
from src.services.report_edit_planner import ReportEditPlanner


class ReportService:
    def __init__(self, edit_planner: ReportEditPlanner | None = None) -> None:
        self._current_report = build_dummy_report()
        self._edit_planner = edit_planner or ReportEditPlanner()

    def get_current_report(self) -> ReportSchema:
        return deepcopy(self._current_report)

    def reset_report(self) -> ReportSchema:
        self._current_report = build_dummy_report()
        return self.get_current_report()

    def set_report(self, report: ReportSchema) -> ReportSchema:
        self._current_report = deepcopy(report)
        return self.get_current_report()

    def create_preview_document(self, report: ReportSchema | None = None) -> Path:
        '''
        Method creates a preview document for the given report or to the current report.
        
        Args:
            report (ReportSchema, optional): The report for which to create the preview document. If None, uses the current report.
        
        Returns:
            pathlib.Path: The path to the created preview document.
        '''
        report_to_render = report or self._current_report
        document = Document()
        self.apply_document_template(document, report_to_render)

        temp_file = NamedTemporaryFile(delete=False, suffix=".docx")
        temp_path = Path(temp_file.name)
        temp_file.close()
        document.save(str(temp_path))
        return temp_path
    
    def apply_document_template(self, document: "Document", report_to_render: ReportSchema) -> None: # type: ignore
        title = document.add_heading(report_to_render.metadata.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        metadata_table = document.add_table(rows=0, cols=2)
        metadata_table.style = "Table Grid"
        self._add_key_value_row(metadata_table, "Report ID", report_to_render.metadata.report_id)
        self._add_key_value_row(metadata_table, "Date", report_to_render.metadata.report_date.isoformat())
        self._add_key_value_row(metadata_table, "Status", report_to_render.metadata.status.title())
        self._add_key_value_row(metadata_table, "Prepared By", report_to_render.metadata.prepared_by)
        self._add_key_value_row(metadata_table, "Client", report_to_render.client.name)
        self._add_key_value_row(metadata_table, "Location", report_to_render.location.address)
        if report_to_render.location.room:
            self._add_key_value_row(metadata_table, "Room", report_to_render.location.room)

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(report_to_render.executive_summary)

        for section in report_to_render.sections:
            document.add_heading(section.heading, level=1)
            document.add_paragraph(section.body)

        document.add_heading("Inspection Issues", level=1)
        for issue in report_to_render.issues:
            self._add_issue(document, issue)

        document.add_heading("Next Steps", level=1)
        for next_step in report_to_render.next_steps:
            document.add_paragraph(next_step, style="List Bullet")


    def apply_requested_changes(self, change_request: ReportChangeRequest) -> tuple[ReportSchema, FilteredReportChange]:
        blocked_change = self._block_obviously_unsafe_request(change_request)
        if blocked_change:
            return self.get_current_report(), blocked_change

        edit_plan = self._edit_planner.create_plan(self._current_report, change_request)
        filtered_change = self._filtered_change_from_plan(edit_plan)

        if edit_plan.intent == "update_report" and edit_plan.patch:
            updated_report, applied_changes, rejected_instructions = self._apply_report_patch(
                self._current_report,
                edit_plan.patch,
            )
            filtered_change.applied_changes.extend(applied_changes)
            filtered_change.rejected_instructions.extend(rejected_instructions)

            if applied_changes:
                self._current_report = updated_report
            elif rejected_instructions:
                filtered_change.intent = "unsupported"
                filtered_change.summary = "The requested edit was not applied because it was unsafe or ambiguous."

        return self.get_current_report(), filtered_change

    def _block_obviously_unsafe_request(self, change_request: ReportChangeRequest) -> FilteredReportChange | None:
        normalized_request = self._normalize_text(change_request.requested_changes)
        security_markers = {
            "api key",
            "bypass safety",
            "developer message",
            "hidden instruction",
            "ignore previous",
            "jailbreak",
            "password",
            "system prompt",
        }
        destructive_actions = {"blank", "clear", "delete", "erase", "remove", "wipe"}
        exact_whole_document_requests = {
            "blank document",
            "blank report",
            "clear document",
            "clear report",
            "delete document",
            "delete report",
            "erase document",
            "erase report",
            "remove document",
            "remove report",
            "wipe document",
            "wipe report",
        }
        tokens = set(normalized_request.split())

        rejected = [marker for marker in security_markers if marker in normalized_request]
        has_destructive_action = bool(destructive_actions.intersection(tokens))
        has_document_target = bool({"document", "report"}.intersection(tokens))
        has_whole_scope = bool({"all", "entire", "everything", "whole"}.intersection(tokens))
        targets_all_content = "content" in tokens or "everything" in tokens
        if (
            normalized_request in exact_whole_document_requests
            or (has_destructive_action and targets_all_content)
            or (has_destructive_action and has_document_target and has_whole_scope)
        ):
            rejected.append("whole document deletion")

        if not rejected:
            return None

        return FilteredReportChange(
            intent="unsupported",
            summary="The request contains instructions outside the safe report-editing scope.",
            rejected_instructions=sorted(set(rejected)),
        )

    @staticmethod
    def _filtered_change_from_plan(edit_plan: ReportEditPlan) -> FilteredReportChange:
        return FilteredReportChange(
            intent=edit_plan.intent,
            summary=edit_plan.summary,
            allowed_instructions=edit_plan.allowed_instructions,
            rejected_instructions=edit_plan.rejected_instructions,
        )

    def _apply_report_patch(
        self,
        report: ReportSchema,
        patch: ReportPatch,
    ) -> tuple[ReportSchema, list[str], list[str]]:
        updated_report = deepcopy(report)
        applied_changes: list[str] = []
        rejected_instructions: list[str] = []

        if patch.metadata:
            applied, rejected = self._apply_model_patch(updated_report.metadata, patch.metadata, "metadata")
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        if patch.client:
            applied, rejected = self._apply_model_patch(
                updated_report.client,
                patch.client,
                "client",
                allowed_none_fields={"email"},
            )
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        if patch.location:
            applied, rejected = self._apply_model_patch(
                updated_report.location,
                patch.location,
                "location",
                allowed_none_fields={"room"},
            )
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        if patch.executive_summary is not None:
            updated_report.executive_summary = patch.executive_summary
            applied_changes.append("Updated executive summary.")

        for section_patch in patch.sections_to_upsert:
            applied, rejected = self._upsert_section(updated_report, section_patch)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        for heading in patch.section_headings_to_remove:
            applied, rejected = self._remove_section(updated_report, heading)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        for issue_patch in patch.issues_to_upsert:
            applied, rejected = self._upsert_issue(updated_report, issue_patch)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        for issue_target in patch.issue_ids_or_titles_to_remove:
            applied, rejected = self._remove_issue(updated_report, issue_target)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        for next_step in patch.next_steps_to_add:
            if self._contains_next_step(updated_report, next_step):
                rejected_instructions.append(f"Next step already exists: {next_step}")
                continue
            updated_report.next_steps.append(next_step)
            applied_changes.append(f"Added next step: {next_step}")

        for next_step_patch in patch.next_steps_to_update:
            applied, rejected = self._update_next_step(updated_report, next_step_patch)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        for next_step in patch.next_steps_to_remove:
            applied, rejected = self._remove_next_step(updated_report, next_step)
            applied_changes.extend(applied)
            rejected_instructions.extend(rejected)

        if not applied_changes:
            return report, [], rejected_instructions

        safety_rejections = self._validate_safe_report_update(report, updated_report)
        if safety_rejections:
            return report, [], rejected_instructions + safety_rejections

        try:
            validated_report = ReportSchema.model_validate(updated_report.model_dump(mode="python"))
        except ValidationError as exc:
            error_message = exc.errors()[0]["msg"] if exc.errors() else str(exc)
            return report, [], rejected_instructions + [f"Patch produced an invalid report: {error_message}"]

        return validated_report, applied_changes, rejected_instructions

    def _apply_model_patch(
        self,
        target_model: Any,
        patch_model: ReportMetadataPatch | ReportPartyPatch | InspectionLocationPatch,
        label: str,
        allowed_none_fields: set[str] | None = None,
    ) -> tuple[list[str], list[str]]:
        allowed_none_fields = allowed_none_fields or set()
        applied_changes: list[str] = []
        rejected_instructions: list[str] = []

        for field_name in patch_model.model_fields_set:
            value = getattr(patch_model, field_name)
            if value is None and field_name not in allowed_none_fields:
                rejected_instructions.append(f"Cannot clear required field: {label}.{field_name}")
                continue

            setattr(target_model, field_name, value)
            applied_changes.append(f"Updated {label}.{field_name}.")

        return applied_changes, rejected_instructions

    def _upsert_section(
        self,
        report: ReportSchema,
        section_patch: ReportSectionPatch,
    ) -> tuple[list[str], list[str]]:
        target_heading = section_patch.match_heading or section_patch.heading
        if target_heading:
            index, match_error = self._find_unique_index(
                [section.heading for section in report.sections],
                target_heading,
            )
            if match_error:
                return [], [f"Section target is ambiguous: {target_heading}"]
            if index is not None:
                section = report.sections[index]
                if section_patch.heading is not None:
                    section.heading = section_patch.heading
                if section_patch.body is not None:
                    section.body = section_patch.body
                return [f"Updated section: {target_heading}"], []

        if section_patch.match_heading:
            return [], [f"Section not found: {section_patch.match_heading}"]

        if not section_patch.heading or not section_patch.body:
            return [], ["New sections require both heading and body."]

        report.sections.append(ReportSection(heading=section_patch.heading, body=section_patch.body))
        return [f"Added section: {section_patch.heading}"], []

    def _remove_section(self, report: ReportSchema, target_heading: str) -> tuple[list[str], list[str]]:
        if self._is_broad_delete_target(target_heading):
            return [], [f"Section deletion target is too broad: {target_heading}"]

        index, match_error = self._find_unique_index(
            [section.heading for section in report.sections],
            target_heading,
        )
        if match_error:
            return [], [f"Section target is ambiguous: {target_heading}"]
        if index is None:
            return [], [f"Section not found: {target_heading}"]

        removed_heading = report.sections.pop(index).heading
        return [f"Removed section: {removed_heading}"], []

    def _upsert_issue(
        self,
        report: ReportSchema,
        issue_patch: ReportIssuePatch,
    ) -> tuple[list[str], list[str]]:
        index, match_error = self._find_issue_index(report, issue_patch)
        if match_error:
            return [], [match_error]

        if index is not None:
            issue = report.issues[index]
            applied_changes: list[str] = []
            rejected_instructions: list[str] = []
            for field_name in ("issue_id", "title", "description", "severity", "recommendation", "image_path"):
                if field_name not in issue_patch.model_fields_set:
                    continue

                value = getattr(issue_patch, field_name)
                if value is None and field_name != "image_path":
                    rejected_instructions.append(f"Cannot clear required issue field: {field_name}")
                    continue

                setattr(issue, field_name, value)
                applied_changes.append(f"Updated issue {issue.issue_id}.{field_name}.")

            return applied_changes, rejected_instructions

        if issue_patch.match_issue_id or issue_patch.match_title:
            target = issue_patch.match_issue_id or issue_patch.match_title
            return [], [f"Issue not found: {target}"]

        missing_fields = [
            field_name
            for field_name in ("issue_id", "title", "description", "severity", "recommendation")
            if getattr(issue_patch, field_name) is None
        ]
        if missing_fields:
            return [], [f"New issues require these fields: {', '.join(missing_fields)}"]

        report.issues.append(
            ReportIssue(
                issue_id=issue_patch.issue_id,
                title=issue_patch.title,
                description=issue_patch.description,
                severity=issue_patch.severity,
                recommendation=issue_patch.recommendation,
                image_path=issue_patch.image_path,
            )
        )
        return [f"Added issue: {issue_patch.issue_id}"], []

    def _remove_issue(self, report: ReportSchema, target: str) -> tuple[list[str], list[str]]:
        if self._is_broad_delete_target(target):
            return [], [f"Issue deletion target is too broad: {target}"]

        index, match_error = self._find_unique_issue_by_target(report, target)
        if match_error:
            return [], [match_error]
        if index is None:
            return [], [f"Issue not found: {target}"]

        removed_issue = report.issues.pop(index)
        return [f"Removed issue: {removed_issue.issue_id}"], []

    def _update_next_step(
        self,
        report: ReportSchema,
        next_step_patch: ReportNextStepPatch,
    ) -> tuple[list[str], list[str]]:
        index, match_error = self._find_unique_index(report.next_steps, next_step_patch.match_text)
        if match_error:
            return [], [f"Next-step target is ambiguous: {next_step_patch.match_text}"]
        if index is None:
            return [], [f"Next step not found: {next_step_patch.match_text}"]

        report.next_steps[index] = next_step_patch.text
        return [f"Updated next step: {next_step_patch.match_text}"], []

    def _remove_next_step(self, report: ReportSchema, target: str) -> tuple[list[str], list[str]]:
        if self._is_broad_delete_target(target):
            return [], [f"Next-step deletion target is too broad: {target}"]

        index, match_error = self._find_unique_index(report.next_steps, target)
        if match_error:
            return [], [f"Next-step target is ambiguous: {target}"]
        if index is None:
            return [], [f"Next step not found: {target}"]

        removed_step = report.next_steps.pop(index)
        return [f"Removed next step: {removed_step}"], []

    def _find_issue_index(
        self,
        report: ReportSchema,
        issue_patch: ReportIssuePatch,
    ) -> tuple[int | None, str | None]:
        if issue_patch.match_issue_id:
            return self._find_unique_index(
                [issue.issue_id for issue in report.issues],
                issue_patch.match_issue_id,
            )
        if issue_patch.match_title:
            return self._find_unique_index(
                [issue.title for issue in report.issues],
                issue_patch.match_title,
            )
        if issue_patch.issue_id:
            index, match_error = self._find_unique_index(
                [issue.issue_id for issue in report.issues],
                issue_patch.issue_id,
            )
            if index is not None or match_error:
                return index, match_error
        if issue_patch.title:
            return self._find_unique_index(
                [issue.title for issue in report.issues],
                issue_patch.title,
            )
        return None, None

    def _find_unique_issue_by_target(self, report: ReportSchema, target: str) -> tuple[int | None, str | None]:
        issue_ids = [issue.issue_id for issue in report.issues]
        index, match_error = self._find_unique_index(issue_ids, target)
        if index is not None or match_error:
            return index, match_error

        issue_titles = [issue.title for issue in report.issues]
        index, match_error = self._find_unique_index(issue_titles, target)
        if match_error:
            return None, f"Issue target is ambiguous: {target}"
        return index, None

    def _find_unique_index(self, candidate_values: list[str], target: str) -> tuple[int | None, str | None]:
        normalized_target = self._normalize_text(target)
        if not normalized_target:
            return None, None

        exact_matches = [
            index
            for index, value in enumerate(candidate_values)
            if self._normalize_text(value) == normalized_target
        ]
        if len(exact_matches) == 1:
            return exact_matches[0], None
        if len(exact_matches) > 1:
            return None, "ambiguous"

        partial_matches = {
            index
            for index, value in enumerate(candidate_values)
            if normalized_target in self._normalize_text(value)
            or self._normalize_text(value) in normalized_target
        }
        if len(partial_matches) == 1:
            return partial_matches.pop(), None
        if len(partial_matches) > 1:
            return None, "ambiguous"

        return None, None

    def _validate_safe_report_update(
        self,
        original_report: ReportSchema,
        updated_report: ReportSchema,
    ) -> list[str]:
        rejections: list[str] = []
        issue_ids = [issue.issue_id for issue in updated_report.issues]

        if updated_report.metadata.report_id != original_report.metadata.report_id:
            rejections.append("Report ID cannot be changed through natural-language edits.")
        if len(issue_ids) != len(set(issue_ids)):
            rejections.append("Issue IDs must remain unique.")
        if not updated_report.sections and not updated_report.issues:
            rejections.append("The report must retain at least one section or inspection issue.")
        if not updated_report.executive_summary.strip():
            rejections.append("The executive summary cannot be empty.")
        if len(updated_report.sections) > 20:
            rejections.append("The report cannot contain more than 20 sections.")
        if len(updated_report.issues) > 30:
            rejections.append("The report cannot contain more than 30 issues.")
        if len(updated_report.next_steps) > 30:
            rejections.append("The report cannot contain more than 30 next steps.")

        return rejections

    def _is_broad_delete_target(self, target: str) -> bool:
        normalized_target = self._normalize_text(target)
        broad_targets = {
            "all",
            "all content",
            "all issues",
            "all next steps",
            "all sections",
            "content",
            "document",
            "entire document",
            "entire report",
            "everything",
            "report",
            "the document",
            "the report",
            "whole document",
            "whole report",
        }
        return normalized_target in broad_targets

    @staticmethod
    def _normalize_text(text: str) -> str:
        return " ".join("".join(character.lower() if character.isalnum() else " " for character in text).split())

    @staticmethod
    def _add_key_value_row(table, key: str, value: str) -> None:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value

    def _add_issue(self, document: Any, issue: ReportIssue) -> None:
        document.add_heading(f"{issue.issue_id}: {issue.title}", level=2)
        document.add_paragraph(f"Severity: {issue.severity.value.title()}")
        document.add_paragraph(issue.description)
        document.add_paragraph(f"Recommendation: {issue.recommendation}")
        if issue.image_path:
            self._add_issue_image(document, issue.image_path)

    def _add_issue_image(self, document: Any, image_path: str) -> None:
        resolved_image_path = self._resolve_image_path(image_path)
        if not resolved_image_path.exists():
            document.add_paragraph(f"Image not found: {image_path}")
            return

        with Image.open(resolved_image_path) as image:
            image.load()
            width, height = image.size

        document.add_picture(str(resolved_image_path), width=Inches(4.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"Image: {resolved_image_path.name} ({width} x {height}px)")

    def _resolve_image_path(self, image_path: str) -> Path:
        path = Path(image_path)
        if path.is_absolute():
            return path
        return _PROJECT_ROOT / path

    @staticmethod
    def _contains_next_step(report: ReportSchema, text: str) -> bool:
        return any(text.lower() in step.lower() for step in report.next_steps)


report_service = ReportService()
